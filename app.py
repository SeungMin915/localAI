import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import plotly.express as px
from datetime import datetime
import os
from docx import Document

# -- 제목 및 소개
st.set_page_config(page_title="HEART Insight AI", layout="wide")
st.title("HEART Insight AI: 현대해상 미래 모빌리티 트렌드 분석 솔루션")

st.markdown("""
본 대시보드는 전기차, 자율주행, MaaS, PBV 등 미래 모빌리티 트렌드를 분석하고  
보험 상품 개발 및 위험 관리를 위한 시사점을 제공합니다.
""")

# -- 데이터 업로드 or 불러오기 (예시 CSV 사용)
st.sidebar.header("데이터 업로드")
uploaded_file = st.sidebar.file_uploader("CSV 파일 업로드", type=["csv"])

if uploaded_file:
    df = pd.read_csv(uploaded_file)
    st.success("데이터 업로드 완료!")
else:
    df = pd.read_csv("sample_data/mobility_trends.csv")  # 예시 데이터
    st.info("예시 데이터를 불러왔습니다.")

# -- 데이터 미리 보기
st.subheader("데이터 미리 보기")
st.dataframe(df.head())

# -- 키워드 기반 트렌드 시각화
st.subheader("트렌드 키워드 시각화 (Plotly)")
if 'keywords' in df.columns:
    keyword_counts = df['keywords'].value_counts().nlargest(10).reset_index()
    keyword_counts.columns = ['Keyword', 'Count']
    fig = px.bar(keyword_counts, x='Keyword', y='Count', title='상위 10개 키워드')
    st.plotly_chart(fig)
else:
    st.warning("키워드 컬럼이 존재하지 않습니다.")

# -- 보험 리스크 레벨 시각화
st.subheader("보험 리스크/기회 수준")
if {'risk_level', 'opportunity_level'}.issubset(df.columns):
    st.write("리스크 수준 분포:")
    st.bar_chart(df['risk_level'].value_counts())

    st.write("기회 수준 분포:")
    st.bar_chart(df['opportunity_level'].value_counts())

# -- 보고서 생성 기능
st.subheader("보고서 생성")
if st.button("보고서 생성 (docx)"):
    now = datetime.now().strftime("%Y%m%d_%H%M%S")
    doc = Document()
    doc.add_heading("HEART Insight AI 분석 보고서", 0)
    doc.add_paragraph(f"보고서 생성일: {now}")

    doc.add_heading("상위 트렌드 키워드", level=1)
    if 'keywords' in df.columns:
        for kw in keyword_counts['Keyword']:
            doc.add_paragraph(f"• {kw}")

    doc.add_heading("분석된 트렌드 수", level=1)
    doc.add_paragraph(f"{len(df)}건")

    filepath = f"HEART_Insight_Report_{now}.docx"
    doc.save(filepath)
    st.success(f"보고서가 생성되었습니다: {filepath}")
    with open(filepath, "rb") as f:
        st.download_button("보고서 다운로드", f, file_name=filepath)

# -- 미래 개발 모듈 안내
with st.expander("🔧 향후 추가 예정 기능"):
    st.markdown("""
    - ARIMA/Prophet/LSTM 기반 시계열 예측
    - 트렌드 간 인과 분석 (Granger Causality)
    - 네트워크 분석 기반 생태계 시각화
    - 자동 리포트 PDF 버전 생성
    - 사용자 맞춤 알림 및 필터링 기능
    """)

